from docx import Document
import string

def reformat_question_docx_to_docx(input_file, output_file):
    document = Document(input_file)
    new_document = Document()
    question_counter = 1  # Стартовый номер вопроса
    question_title = ""  # Accumulate question title across multiple lines
    answer_index = 0  # Reset answer index for each question
    current_question = None  # Keep track of the current question

    for para in document.paragraphs:
        line = para.text.strip()
        if line.startswith(('V1', 'V2')):
            if current_question:
                # Add the previous question to the new document
                new_document.add_paragraph(f"{current_question['number']}. {current_question['title']}")
                for answer in current_question['answers']:
                    new_document.add_paragraph(f"{answer['letter']}\t{answer['text']}")

            # Reset the current question
            current_question = {
                'number': question_counter,
                'title': line.split('\t')[1].strip(),
                'answers': []
            }
            question_counter += 1
            answer_index = 0
        elif line.startswith(('0', '1')):
            if not current_question:
                # Skip this answer if we're not in a question
                continue

            answer = line[2:].strip()
            letter = string.ascii_uppercase[answer_index]
            current_question['answers'].append({'letter': letter, 'text': answer})
            answer_index += 1
        else:
            if current_question:
                # Accumulate the question title across multiple lines
                current_question['title'] += line + " "

    if current_question:
        # Add the last question to the new document
        new_document.add_paragraph(f"{current_question['number']}. {current_question['title']}")
        for answer in current_question['answers']:
            new_document.add_paragraph(f"{answer['letter']}\t{answer['text']}")

    new_document.save(output_file)

reformat_question_docx_to_docx('тест.docx', 'output.docx')